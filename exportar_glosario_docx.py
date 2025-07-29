from docx import Document
from docx.shared import Pt
import json

# 1. Cargar el glosario enriquecido
with open("glosario_senado.json", "r", encoding="utf-8") as file:
    glosario = json.load(file)

# 2. Crear documento Word
doc = Document()
doc.add_heading("Glosario Señado Institucional", 0)

# 3. Configurar estilo básico
style = doc.styles['Normal']
font = style.font
font.name = 'Segoe UI'
font.size = Pt(11)

# 4. Escribir cada entrada
for palabra_data in glosario:
    palabra = palabra_data["palabra"]
    descripcion = palabra_data["descripcion"]
    seña = palabra_data["seña"]

    doc.add_heading(f"📌 {palabra.capitalize()}", level=2)
    doc.add_paragraph(f"🧭 {descripcion}", style='Normal')
    doc.add_paragraph(f"✋ {seña}", style='Normal')
    doc.add_paragraph("")  # Espacio entre entradas

# 5. Guardar el documento
doc.save("Glosario Señado Institucional.docx")
print("✅ Documento Word generado exitosamente como 'Glosario Señado Institucional.docx'")